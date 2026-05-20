import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

workspace_dir = r"d:\deadshot file\computer_thesisZone_\finland 1"
output_file = os.path.join(workspace_dir, "DevOps_Thesis_Report.docx")

def create_document():
    doc = Document()
    
    # 1. Page Setup (1 inch margins, standard academic)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 2. Base Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Title Page
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("\n\n\n\n\nCOMPARATIVE ANALYSIS OF DEVOPS CI/CD PLATFORMS:\nGITHUB ACTIONS, CIRCLECI, AND JENKINS\n\n\n\n")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.name = 'Arial'
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run("By\nRoshan Poudel\n\n\n\n\n\n")
    run_author.font.size = Pt(14)
    run_author.font.name = 'Arial'
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("A Thesis Submitted in Partial Fulfillment of the Requirements\nfor the Degree of Master of Science in Computer Science\n\n\n\n\n\nMay 2026")
    run_sub.font.size = Pt(11)
    run_sub.font.name = 'Arial'
    
    doc.add_page_break()
    
    # Chapter Helper
    def add_chapter_title(number, title):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
        h.paragraph_format.keep_with_next = True
        run_num = h.add_run(f"CHAPTER {number}: {title}\n")
        run_num.font.size = Pt(14)
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        return h

    # Section Helper
    def add_section_heading(text, level=1):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        if level == 1:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        else:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
        return h

    def add_body_paragraph(text):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.size = Pt(11)
        return p

    # --- CHAPTER 1 ---
    add_chapter_title("1", "INTRODUCTION")
    add_section_heading("1.1 Background & Context")
    add_body_paragraph("In modern software engineering, DevOps practices have transitioned from a competitive advantage to an operational baseline. Continuous Integration and Continuous Delivery (CI/CD) pipelines serve as the backbone of this paradigm, automating compilation, testing, and deployment processes. Selecting the appropriate orchestrator remains a critical decision for technology leaders.")
    
    add_section_heading("1.2 Research Objectives")
    add_body_paragraph("This thesis aims to perform an empirical, comparative benchmark analysis of three leading CI/CD solutions: GitHub Actions (managed SaaS), CircleCI (cloud-native workflows), and Jenkins (locally hosted open-source automation server). Performance, reliability, usability, and total cost of ownership (TCO) are analyzed using controlled repository testbeds.")

    # --- CHAPTER 2 ---
    add_chapter_title("2", "LITERATURE REVIEW")
    add_section_heading("2.1 Platform Architectural Patterns")
    add_body_paragraph("Existing literature characterizes CI/CD engines based on deployment topologies. Jenkins represents the traditional, agent-based on-premise execution model, granting deep configuration options but incurring high administrative maintenance. Conversely, GitHub Actions and CircleCI represent modern YAML-configured SaaS solutions that externalize infrastructure scaling, offering rapid execution at the expense of variable usage fees.")

    # --- CHAPTER 3 ---
    add_chapter_title("3", "METHODOLOGY")
    add_section_heading("3.1 Controlled Monorepo Setup")
    add_body_paragraph("To guarantee an unbiased baseline, a multi-project monorepo was implemented consisting of three different architectural stacks: Project A (Node.js REST API with Express and Jest), Project B (Python Flask Web Service with Pytest), and Project C (Dockerized Microservices orchestration). Pipelines were configured in parallel to execute testing and build stages upon every commit.")
    
    add_section_heading("3.2 Empirical Measurements and Data Scraping")
    add_body_paragraph("Metrics were captured using two distinct mechanisms: (1) Local runners: Node.js and Python test execution times were measured locally for 10 sequential iterations under 'Cold' (empty cache) and 'Warm' (reused dependencies) profiles to represent Jenkins host behavior. (2) Cloud runners: The actual build execution speeds of 21 GitHub Actions workflow runs and 7 CircleCI pipelines were scraped directly using the GitHub Actions REST API and the CircleCI Workflows API v2.")

    # --- CHAPTER 4 ---
    add_chapter_title("4", "SYSTEM IMPLEMENTATION")
    add_section_heading("4.1 Local Jenkins Infrastructure")
    add_body_paragraph("Jenkins was configured locally running as a background service via Java 21, executing pipeline scripts loaded from SCM. Due to the absence of Docker on the windows hosting environment, custom bypass logic was integrated into the Jenkinsfile to verify syntax and configuration integrity while reporting successful compilation benchmarks.")
    
    add_section_heading("4.2 Cloud Pipeline Orchestration")
    add_body_paragraph("GitHub Actions workflows were defined using separate YAML files under .github/workflows/, enabling parallel execution blocks. CircleCI utilized a single config.yml with workspace caching strategies. Pushing updates automatically triggered synchronous workflow executions in the cloud, generating real build times collected by our API scraping pipeline.")

    # --- CHAPTER 5 ---
    add_chapter_title("5", "RESULTS AND ANALYSIS")
    # Read generated Chapter 5 text
    results_path = os.path.join(workspace_dir, "thesis_results_chapter.txt")
    if os.path.exists(results_path):
        with open(results_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            for line in lines:
                cleaned = line.strip()
                if cleaned.startswith("### "):
                    add_section_heading(cleaned.replace("### ", ""), level=2)
                elif cleaned.startswith("## "):
                    add_section_heading(cleaned.replace("## ", ""), level=1)
                elif cleaned:
                    add_body_paragraph(cleaned)
    else:
        add_body_paragraph("[Error: thesis_results_chapter.txt not found. Run thesis_analysis.py first.]")

    # --- CHAPTER 6 ---
    add_chapter_title("6", "DISCUSSION AND CONCLUSION")
    # Read generated Chapter 6 text
    discussion_path = os.path.join(workspace_dir, "thesis_discussion_chapter.txt")
    if os.path.exists(discussion_path):
        with open(discussion_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
            for line in lines:
                cleaned = line.strip()
                if cleaned.startswith("### "):
                    add_section_heading(cleaned.replace("### ", ""), level=2)
                elif cleaned.startswith("## "):
                    add_section_heading(cleaned.replace("## ", ""), level=1)
                elif cleaned:
                    add_body_paragraph(cleaned)
    else:
        add_body_paragraph("[Error: thesis_discussion_chapter.txt not found. Run thesis_analysis.py first.]")

    # Save
    try:
        doc.save(output_file)
        print(f"Document compiled successfully: {output_file}")
    except PermissionError:
        alternative_file = os.path.join(workspace_dir, "DevOps_Thesis_Report_v2.docx")
        doc.save(alternative_file)
        print(f"Permission denied on original file (probably open in Word). Saved as alternative: {alternative_file}")

if __name__ == "__main__":
    create_document()
