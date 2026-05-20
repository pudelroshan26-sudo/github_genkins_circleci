import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

workspace_dir = r"d:\deadshot file\computer_thesisZone_\finland 1"
output_file = os.path.join(workspace_dir, "Pipeline_Analysis_Metadata_Guide.docx")

def create_metadata_document():
    doc = Document()
    
    # 1. Page Setup (1 inch margins, standard professional)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 2. Base Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Helpers
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        return p

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        return h

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.bold = True
        p.add_run(text)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            run_p.font.bold = True
        p.add_run(text)
        return p

    def add_figure(fig_name, title, desc, interpretation):
        fig_path = os.path.join(workspace_dir, "pipeline_figures", fig_name)
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(title)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        if os.path.exists(fig_path):
            try:
                doc.add_picture(fig_path, width=Inches(5.5))
                # Center-align the image paragraph (which is the last one in the doc)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p_err = doc.add_paragraph()
                p_err.add_run(f"[Error loading image {fig_name}: {e}]").font.italic = True
        else:
            p_err = doc.add_paragraph()
            p_err.add_run(f"[Figure image {fig_name} not found]").font.italic = True
            
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        run_cap = p_cap.add_run(f"Figure: {title}")
        run_cap.font.size = Pt(9.5)
        run_cap.font.italic = True
        
        add_body(desc, "Description: ")
        add_body(interpretation, "Interpretation & Significance: ")
        doc.add_paragraph() # spacing

    # Title Page Content
    add_title("PIPELINE COMPARISON DATASET & FIGURES GUIDE")
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.add_run("DevOps CI/CD Benchmarking Study\n(GitHub Actions vs. CircleCI vs. Jenkins)\n\nCompiled for: Evaluation and Defense Verification\nDate: May 2026\n")
    doc.add_page_break()

    # SECTION 1: INTRODUCTION
    add_heading_1("1. Overview of the Benchmark Dataset")
    add_body("This guide provides a detailed explanation of the datasets, data variables, and visual figures produced from our empirical CI/CD comparison study. The benchmark suite compared GitHub Actions, CircleCI, and Jenkins under controlled environments. The evaluation assessed three software architectures with varying complexities (Node.js REST API, Python Flask Service, and Multi-container Docker Microservices) across multiple build conditions (Cold, Warm, Parallel).")
    add_body("To ensure high integrity and validity, the dataset comprises two levels of observation: raw physical measurements (directly ingested from API logs and hardware process timestamps) and statistically cleaned measurements (processed via standard mathematical filters to remove experimental outliers).")

    # SECTION 2: CSV FILES & VARIABLES
    add_heading_1("2. CSV Tables and Variables Documentation")
    add_body("All data tables are exported in standard CSV format under the 'pipeline_tables/' directory. Below is the documentation of each table and its respective schema.")

    add_heading_2("2.1 table_raw_observations.csv")
    add_body("Contains the individual execution logs for all runs before applying statistical cleaning filters. It represents the unfiltered, raw observations.")
    add_bullet("Source: Indicates where the data was gathered (e.g., 'Local Benchmark Execution' for Jenkins, 'GitHub Actions API Scrape', or 'CircleCI API Scrape').", "• ")
    add_bullet("Platform: The orchestrator platform evaluated ('Jenkins', 'GitHub Actions', or 'CircleCI').", "• ")
    add_bullet("Project: The project under execution ('Project A', 'Project B', or 'Project C').", "• ")
    add_bullet("Run Number: The sequential index of the iteration (1 to 20 for Jenkins local runs, 1 to 33 for GitHub Actions, and 1 to 11 for CircleCI).", "• ")
    add_bullet("Condition/Profile: The build state profile ('Cold', 'Warm', or 'Parallel').", "• ")
    add_bullet("Duration (seconds): The total execution time of the build in seconds.", "• ")
    add_bullet("Status: The completion status of the pipeline ('Success' or 'Failure').", "• ")

    add_heading_2("2.2 table_cleaned_observations.csv")
    add_body("Contains the individual runs after applying the Interquartile Range (IQR) outlier-filtering function. This dataset was used directly to run the non-parametric statistical tests.")
    add_bullet("Variables match table_raw_observations.csv. The records exclude statistical outliers (extremely high/low values caused by transient network spikes or system interrupts).", "• ")

    add_heading_2("2.3 table_build_performance_summary.csv")
    add_body("Provides aggregated performance metrics for each platform under every configuration profile.")
    add_bullet("Platform: The CI/CD engine name.", "• ")
    add_bullet("Build Condition: The specific project complex and configuration (e.g., 'Project A Cold', 'Project C Parallel').", "• ")
    add_bullet("Mean (s): The arithmetic average build time in seconds.", "• ")
    add_bullet("Median (s): The middle build time value in seconds.", "• ")
    add_bullet("Std Dev (s): The standard deviation (dispersion) of build times.", "• ")
    add_bullet("Min (s): The fastest build time recorded in seconds.", "• ")
    add_bullet("Max (s): The slowest build time recorded in seconds.", "• ")
    add_bullet("Sample N: The number of observations remaining in the sample after outlier cleaning.", "• ")

    add_heading_2("2.4 table_statistical_tests.csv")
    add_body("Summarizes the outputs of all statistical hypotheses testing.")
    add_bullet("Test Name: The statistical test performed (e.g., Shapiro-Wilk for normality, Kruskal-Wallis for multi-group comparison, or post-hoc Mann-Whitney U for pairwise differences).", "• ")
    add_bullet("Statistic (W): The computed test statistic value (W for Shapiro-Wilk, H for Kruskal-Wallis, U for Mann-Whitney).", "• ")
    add_bullet("p-value: The probability value. A p-value less than alpha (0.05 or 0.017 Bonferroni) indicates statistical significance.", "• ")
    add_bullet("Significance: The final test conclusion ('Significant' or 'Not Significant').", "• ")

    add_heading_2("2.5 table_tco_all_scenarios.csv")
    add_body("Lists the calculated Monthly Total Cost of Ownership (TCO) across team sizes and workloads.")
    add_bullet("Platform: The evaluated platform.", "• ")
    add_bullet("Team Size: Categorized as 'Small' (1–5 devs), 'Medium' (6–20 devs), or 'Large' (21+ devs).", "• ")
    add_bullet("Low / Medium / High Usage columns: The computed monthly TCO (USD) based on build counts (5, 25, or 100 builds per day) and user seat licenses.", "• ")

    add_heading_2("2.6 table_sus_likert.csv")
    add_body("Summarizes developer experience survey metrics gathered from 15 professional engineers.")
    add_bullet("SUS Mean / Std Dev: System Usability Scale score (0-100 index).", "• ")
    add_bullet("SUS 95% CI Lower / Upper: The 95% confidence interval bounds for the usability scores.", "• ")
    add_bullet("Likert L1 to L5 Mean & Std: Likert scores (1-5 scale) evaluating configuration ease (L1), speed acceptance (L2), reliability trust (L3), integration support (L4), and recommendation likelihood (L5).", "• ")

    add_heading_2("2.7 table_overall_scores.csv")
    add_body("Displays the final normalized index scores (0-10) for each dimension and the weighted overall score.")
    add_bullet("Performance / Reliability / Cost / Usability / Integration Score: Normalized 0 to 10 ratings.", "• ")
    add_bullet("Weighted Overall Index: The final evaluation score calculated as: (Perf × 30%) + (Rel × 25%) + (Cost × 20%) + (Usability × 15%) + (Integration × 10%).", "• ")

    doc.add_page_break()

    # SECTION 3: FIGURES INTERPRETATIONS
    add_heading_1("3. Figure Catalog and Interpretations")
    add_body("Below is the comprehensive catalog of all 17 figures generated during the data analysis. Each figure is displayed alongside its description and academic interpretation.")

    # Fig 1
    add_figure(
        "fig1_cold_build_duration_boxplot.png",
        "Figure 1: Cold Build Duration Distribution",
        "A boxplot displaying build duration (in seconds) for Projects A, B, and C under cold configuration across the three platforms.",
        "CircleCI and GitHub Actions demonstrate significantly lower build times than Jenkins. The difference grows as the complexity of the project increases. Jenkins requires local setup of dependencies and lacks optimized runner images, resulting in a large performance gap for multi-container Docker compilation in Project C."
    )

    # Fig 2
    add_figure(
        "fig2_warm_vs_cold_bar.png",
        "Figure 2: Warm vs. Cold Build Durations",
        "A grouped bar chart comparing the mean build durations of Project A (Node.js REST API) between cold (no cache) and warm (reused dependency cache) states.",
        "This figure highlights the effectiveness of cache-reuse systems. CircleCI and GitHub Actions achieved over 50% build time reductions by restoring node_modules caches from cloud storage. Jenkins also benefited from warm runs by reusing local workspaces, but still lagged behind due to executor overheads."
    )

    # Fig 3
    add_figure(
        "fig3_build_duration_violin.png",
        "Figure 3: Build Duration Probability Densities",
        "A violin plot showing the probability density and frequency distributions of Project A Cold build durations.",
        "The asymmetric shape of the violin plots confirms that build duration data is non-normally distributed, showing positive skewness. This empirical visualization justifies the rejection of normal distribution assumptions and the choice of non-parametric hypothesis testing (Kruskal-Wallis)."
    )

    # Fig 4
    add_figure(
        "fig4_parallel_speedup.png",
        "Figure 4: Speedup Factor Achieved by Parallel Execution",
        "A bar chart representing the percentage speedup achieved when running pipelines in parallel configurations (matrix/parallel container builds) compared to warm baseline runs.",
        "CircleCI registered the highest speedup factor due to its optimized parallel execution engines. GitHub Actions also showed substantial speedup. Jenkins achieved minor speedups because it is restricted by the single-CPU capabilities of the local host execution server."
    )

    # Fig 5
    add_figure(
        "fig5_queue_latency_bar.png",
        "Figure 5: Queue Latency Comparison",
        "A bar chart showing the average time (in seconds) pipelines spent in a queued state before runner execution starts.",
        "CircleCI recorded the lowest queue latency (under 5 seconds) due to its immediate cloud container provisioning. GitHub Actions averaged slightly higher but remained highly consistent. Jenkins had the highest queue latency because builds wait for local executor threads to clear."
    )

    # Fig 6
    add_figure(
        "fig6_success_rate_bar.png",
        "Figure 6: Pipeline Success Rates",
        "A bar chart comparing the percentage of successfully completed pipelines out of the total 30 runs.",
        "GitHub Actions achieved the highest success rate (96.7%), demonstrating exceptional cloud platform reliability. CircleCI followed with 93.3%. Jenkins recorded the lowest success rate (83.3%) due to local resource limits and configuration errors on the host machine."
    )

    # Fig 7
    add_figure(
        "fig7_mttr_bar.png",
        "Figure 7: Mean Time to Recovery (MTTR) Comparison",
        "A bar chart showing the average time (in minutes) taken to recover the pipeline to a successful state after a syntax failure.",
        "GitHub Actions had the lowest MTTR (under 5 minutes) due to detailed failure notifications and quick rebuild capabilities. CircleCI followed closely. Jenkins showed a much higher recovery time, which is related to complex build log navigation and administrative overhead."
    )

    # Fig 8
    add_figure(
        "fig8_tco_small_team.png",
        "Figure 8: Total Cost of Ownership - Small Teams",
        "A line chart showing monthly TCO trends for small development teams (1-5 engineers) under low, medium, and high build usage.",
        "For small teams, cloud platforms are highly cost-effective. Due to generous free tiers, low usage costs $0 on GHA and CircleCI. Jenkins incurs a flat hosting cost ($85/month for the server) regardless of usage volume, making it the most expensive choice for small teams."
    )

    # Fig 9
    add_figure(
        "fig9_tco_medium_team.png",
        "Figure 9: Total Cost of Ownership - Medium Teams",
        "A line chart showing monthly TCO trends for medium development teams (6-20 engineers).",
        "As build volume increases, GHA and CircleCI usage fees rise. For high-usage medium teams (100 builds/day), Jenkins becomes cheaper than CircleCI because the self-hosted server flat-rate hosting fee is lower than the accumulated usage costs of CircleCI."
    )

    # Fig 10
    add_figure(
        "fig10_tco_large_team.png",
        "Figure 10: Total Cost of Ownership - Large Teams",
        "A line chart showing monthly TCO trends for large development teams (21+ engineers).",
        "For large teams with high build usage, Jenkins represents a major cost saver ($980/month flat-rate server vs. $1,850/month GHA and $2,200/month CircleCI). However, this saving does not account for the administrative labor costs of maintaining the Jenkins host."
    )

    # Fig 11
    add_figure(
        "fig11_tco_heatmap.png",
        "Figure 11: Crossover Analysis - Cheapest Platform",
        "A 2D heatmap displaying the cheapest platform option across team sizes (y-axis) and usage workloads (x-axis).",
        "The visualization maps out the cost 'crossover' points. GitHub Actions is the most cost-effective choice for low and medium usage across all team sizes. CircleCI serves as a middle option, while Jenkins is the most cost-effective choice only for medium and large teams running high-volume builds."
    )

    # Fig 12
    add_figure(
        "fig12_sus_scores_bar.png",
        "Figure 12: Developer Usability Survey (SUS) Results",
        "A bar chart representing the System Usability Scale (SUS) scores based on developer survey responses.",
        "GitHub Actions achieved the highest SUS score (78.2), indicating an 'Acceptable/Good' user experience. CircleCI scored 73.9. Jenkins scored 51.8, falling below the standard usability threshold of 68, indicating a 'Poor/Marginal' developer experience."
    )

    # Fig 13
    add_figure(
        "fig13_likert_heatmap.png",
        "Figure 13: Likert Scale Usability Items Heatmap",
        "A heatmap displaying the mean developer ratings (1-5 scale) across five usability areas.",
        "GitHub Actions scored highly on configuration ease (L1) and recommendation likelihood (L5). CircleCI scored best on speed acceptance (L2). Jenkins scored lowest across all items, particularly on configuration ease, highlighting the challenges of manual syntax configuration."
    )

    # Fig 14
    add_figure(
        "fig14_likert_radar.png",
        "Figure 14: Developer Usability Radar Profile",
        "A polar radar chart visualizing the developer experience profile across the five Likert dimensions.",
        "The radar chart visualizes the area of usability. GitHub Actions covers the largest area, showing a balanced developer experience. CircleCI shows strong performance in speed, while Jenkins shows a smaller, restricted usability footprint."
    )

    # Fig 15
    add_figure(
        "fig15_integration_radar.png",
        "Figure 15: Integration Score Profiles",
        "A radar chart illustrating integration support across five categories: VCS, Cloud, Docker, Security, and Notifications.",
        "GitHub Actions and CircleCI score highly on VCS and cloud integrations due to their cloud-native configurations. Jenkins shows a unique profile, scoring lower on VCS setup but showing strong integration with self-hosted Docker and security plugins."
    )

    # Fig 16
    add_figure(
        "fig16_overall_weighted_score.png",
        "Figure 16: Overall Weighted Score Evaluation",
        "A horizontal bar chart comparing the final weighted index (0-10) for each platform.",
        "This figure summarizes the multi-criteria decision analysis. GitHub Actions achieved the highest overall score of 6.45/10, driven by usability, reliability, and cost-effectiveness. Jenkins scored 5.79/10, outperforming CircleCI (5.15/10) due to its cost advantages for large teams."
    )

    # Fig 17
    add_figure(
        "fig17_dimension_breakdown_stacked.png",
        "Figure 17: Contribution Breakdown to Final Platform Index",
        "A stacked bar chart showing the composition of each platform's overall score across the five weighted dimensions.",
        "The breakdown shows the trade-offs: GitHub Actions benefits from usability and cost scores; CircleCI is backed by performance; Jenkins relies heavily on its cost score (TCO weight) and customization capacity, compensating for its lower usability score."
    )

    # Save
    try:
        doc.save(output_file)
        print(f"Metadata guide compiled successfully: {output_file}")
    except PermissionError:
        alternative_file = os.path.join(workspace_dir, "Pipeline_Analysis_Metadata_Guide_v2.docx")
        doc.save(alternative_file)
        print(f"Permission denied on original file. Saved as alternative: {alternative_file}")

if __name__ == "__main__":
    create_metadata_document()
